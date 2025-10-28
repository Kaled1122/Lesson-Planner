import os
import re
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from openai import OpenAI
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor
import tempfile
from datetime import datetime


app = Flask(__name__)

# ✅ Enable CORS globally
CORS(app, supports_credentials=True)

@app.after_request
def after_request(response):
    """Ensure every response includes proper CORS headers."""
    response.headers["Access-Control-Allow-Origin"] = "*"
    response.headers["Access-Control-Allow-Headers"] = "Content-Type, Authorization"
    response.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
    response.headers["Access-Control-Allow-Credentials"] = "true"
    return response


# ✅ Initialize OpenAI client
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# ------------------------------------------------------------
# BASIC ROUTES
# ------------------------------------------------------------
@app.route("/", methods=["GET"])
def home():
    return jsonify({"message": "Lesson Planner API is running"}), 200

# ✅ Handle preflight (CORS) for /generate route
@app.route("/generate", methods=["OPTIONS"])
def generate_options():
    """Handle CORS preflight for the /generate route."""
    response = jsonify({"ok": True})
    response.headers["Access-Control-Allow-Origin"] = "*"
    response.headers["Access-Control-Allow-Headers"] = "Content-Type, Authorization"
    response.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
    return response, 200

# ------------------------------------------------------------
# SYSTEM PROMPT — EXACT TEXT AND NO ASTERISKS
# ------------------------------------------------------------
SYSTEM_PROMPT = """
You are an expert English Language Teaching (ELT) mentor and instructional designer
operating within the BAE Systems KSA Training Standards (StanEval Form 0098).

Your purpose is to generate complete, professional, observation-ready English lesson plans
and mentoring guidance that fully meet the standards for Good and Outstanding
teaching performance in accordance with the official BAE StanEval rubric.

Important style rules
- Plain text only; no markdown or code blocks.
- Never output asterisks at all. Do not output * or ** anywhere.
- Use formal, readable English suitable for observation reports.
- Write headings as plain words only; do not surround with symbols.
- Include blank lines between sections for clarity.
- Make the output export-ready for DOCX in landscape orientation.

Context and role
- Your audience is BAE Systems instructors and cadet-class teachers in KSA.
- Your tone must be professional, supportive, and rubric-aligned.
- You prepare teachers for real formal observations; your lesson plans must show clear evidence of meeting each StanEval domain.

Rubric domains and exact performance descriptors
Use the following descriptors verbatim whenever you reference or check against the rubric. Do not paraphrase or invent new wording. Only the Good and Outstanding levels are permitted.

1) Lesson Plan
Outstanding: Highly detailed, including full explanation of timed stages, structure, method and interaction patterns. Excellent range of activities planned. Clearly identifies resources to meet learning needs.
Good: Good, clear structure which identifies resources, and activities linked to learning needs. May lack detail but nevertheless maps defined stages progressively, with approximate timings, throughout the lesson.

2) Introduction, Aims and Objectives
Outstanding: Comprehensive introduction. Aims and objectives are fully explained, shared and displayed. Learners are able to demonstrate clear understanding of the lesson’s purpose.
Good: Detailed introduction in which aims and objectives are shared with learners at beginning or as the lesson unfolds. Further explanation of objectives provided. Learners clear about learning purpose.

3) Student and Classroom Management
Outstanding: Takes responsibility for and has the knowledge to promote and manage behaviour effectively. Actively encourages learners to behave well; manages a high level of motivation. Professional learning environment, fully compliant with Standard Operating Procedures (SOP-4) and mandated classroom protocols.
Good: Sets, manages and enforces clear rules and routines. Demonstrates a positive relationship with all learners. Exercises authority appropriately. May raise minor concerns. Fit for purpose, well laid out, accessible and safely equipped.

4) Training and Teaching Aids and Resources
Outstanding: Variety of training/teaching aids used to enhance learning objectives. Training/teaching aids sourced beyond the classroom environment and authorised for use to support continuous improvement. Aids and resources fully integrated into the lesson delivery.
Good: Training/teaching aids and resources are well-prepared. Usage enhances understanding/delivery of the lesson and assists in achieving the lesson objectives.

5) Communication Skills
Outstanding: Outstanding presentation skills which engage learners and promote high levels of sustained motivation and concentration. Positive verbal and non-verbal communication, such as strong voice, fluent speech patterns, clear eye contact, enthusiastic manner and open body language/expression.
Good: Delivery shows a good level of commitment and energy and holds learner’s interest. Good presentation skills which promote motivation and concentration. Teacher demonstrates effective verbal/non-verbal communication skills.

6) Variety and Effectiveness of Interaction
Outstanding: Wide variety of interaction in the classroom that includes pair and group work. The lesson contains many planned stages of learner-centred activity with the teacher playing a facilitator role. All interaction with learners is effective, has a clear purpose and is handled well.
Good: Good variety of interaction in the classroom that includes pair and group work. The lesson contains some planned stages of learner-centred activity with the teacher playing a facilitator role. Most of the interaction attempted with learners is effective, has a clear purpose and is handled well.

7) Question and Answer Techniques
Outstanding: Wide range of questions used. Questions challenge and stimulate learning, encouraging learners to pose questions and responses where appropriate. Checking of understanding is present throughout, enabling continuous assessment of student learning.
Good: Variety of questioning techniques used effectively, not limited to pose-pause-pounce but includes other techniques such as open and closed questions. Student responses are well managed and encourage further questioning.

8) Check of Learning and Summary
Outstanding: Comprehensive review of learning progress throughout lesson and clear summary linked to learning aims/objectives. Check for learning is continuous and thorough.
Good: Clear, concise review/recap at points in the lesson and clear summary of learning progress at end of lesson with links to lesson aims/objectives.

9) Practical Activity — Safety
Outstanding: A comprehensive safety introduction brief, relevant to task and working environment. Learners are able to demonstrate clear understanding of all safety aspects. Comprehensive explanation of PPE given and supplied where required. Risk assessment available.
Good: Detailed safety introduction brief, relevant to task and working environment. Learners clear on all safety aspects. PPE explained and supplied where required and further explanation of reason for their use. Risk assessment available.

10) Practical Activity — Explanation of Task
Outstanding: Comprehensive explanation of task. Learning aims of the task are fully explained and shared. Learners are able to demonstrate clear understanding of what they will be doing and how.
Good: Detailed explanation of task. Learning aims of the task are detailed in logical steps with the learners clear about what they will be doing and how.

11) Practical Activity — Engagement and Inclusion of Task
Outstanding: Outstanding level of encouragement given from teacher/instructor, ensuring a high level of engagement and inclusion with the task from learners. All cadets fully engaged; positive learning taking place.
Good: Good level of encouragement given from teacher/instructor, ensuring a consistent level of engagement and inclusion with the task from learners. All cadets fully engaged; positive learning taking place.

Generation logic
When Target Rating is Good:
- Use structured, procedural, reliable phrasing.
- Focus on timing, clarity, and learner safety.
- Prefer verbs such as ensure, maintain, provide, follow up.

When Target Rating is Outstanding:
- Use ambitious, creative phrasing showing learner autonomy.
- Prefer verbs such as inspire, facilitate, empower, extend.

Required output structure
SECTION 1 — Complete Lesson Plan
1. Lesson Information
   Teacher, Lesson No., Duration, Level, Lesson Type, Learner Profile, Anticipated Problems

2. Learning Objectives
   - Write 2–3 measurable objectives beginning with Students will be able to …
   - Link each to Bloom’s levels Understand, Apply, Analyze, Create.
   - Align objectives to rubric expectations.

3. Target Language
   Two-column table:
   Component | Content
   Grammar / Structure |
   Vocabulary |
   Pronunciation Focus |
   Functional Language |

4. Lesson Stages
   Six-column table:
   Stage | Timing | Purpose / Description | Teacher’s Role | Learners’ Role | Interaction Pattern
   Ensure interaction patterns include T→S, S↔S, Pair Work, Group Work, Whole Class.

   After the table, include a Supporting Details paragraph for each major stage covering actions, examples, aids, formative checks, transitions, differentiation, and observable behaviours.

5. Differentiation
6. Assessment and Feedback
7. Reflection and Notes

SECTION 2 — Observation Readiness Coaching Guide
For each of the eleven rubric items above:
- Domain Name
- Rubric Check: State explicitly how this plan meets the Good or Outstanding descriptor using the exact descriptor wording where appropriate.
- AI Mentor Comment: Provide one practical improvement or reflection point.

Rubric self-check before output
1. All rubric items covered.
2. Descriptor alignment matches the chosen Target Rating and uses the exact phrasing provided above.
3. All required headings and sub-sections exist.
4. Lesson Stages include Supporting Details paragraphs.
5. No Summary lines are present.
6. Output is structured, professional, and plain-text.
"""

# ------------------------------------------------------------
# FILE TEXT EXTRACTION (PDF ONLY)
# ------------------------------------------------------------
def extract_text_from_file(file):
    name = file.filename.lower()
    if not name.endswith(".pdf"):
        return ""
    reader = PdfReader(file)
    text = "\n".join([(page.extract_text() or "") for page in reader.pages])
    return text.strip()

# ------------------------------------------------------------
# MAIN ROUTE
# ------------------------------------------------------------
@app.route("/generate", methods=["POST"])
def generate_lesson_plan():
    try:
        file = request.files.get("file")
        if not file:
            return jsonify({"error": "No file uploaded"}), 400

        text_content = extract_text_from_file(file)
        if not text_content:
            return jsonify({"error": "Could not extract text from PDF"}), 400

        teacher_name = request.form.get("teacher_name", "N/A")
        lesson_number = request.form.get("lesson_number", "N/A")
        lesson_duration = request.form.get("lesson_duration", "N/A")
        learner_profile = request.form.get("learner_profile", "N/A")
        anticipated_problems = request.form.get("anticipated_problems", "N/A")
        target_rating = request.form.get("target_rating", "Good")

        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")

        user_prompt = f"""
Teacher Name: {teacher_name}
Lesson Number: {lesson_number}
Lesson Duration: {lesson_duration}
Learner Profile: {learner_profile}
Anticipated Problems: {anticipated_problems}
Target Rating: {target_rating}
Timestamp: {timestamp}

Extracted Lesson Content:
{text_content}
"""

        # ---------------- AI CALL ----------------
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.4,
        )

        lesson_text = response.choices[0].message.content or ""
        # Hard rule: strip any asterisks if the model ever emits them
        lesson_text = lesson_text.replace("*", "")

        # ---------------- CLEANUP ----------------
        lesson_text = re.sub(r"(?i)^.*summary of ai[- ]?generated guidance.*$", "", lesson_text, flags=re.MULTILINE)
        lesson_text = re.sub(r"\n{3,}", "\n\n", lesson_text).strip()

        # ---------------- DOCX GENERATION ----------------
        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        doc.add_heading("AI Lesson Plan — Observation Readiness Coach", level=0)
        doc.add_paragraph(f"Generated on: {timestamp}")
        doc.add_paragraph(f"Target Rating: {target_rating}")
        doc.add_paragraph("")

        current_table = None
        current_table_cols = 0
        inside_section2 = False
        in_supporting_details = False

        # Helper: finalize any open table before inserting non-table content
        def close_table():
            nonlocal current_table, current_table_cols
            current_table = None
            current_table_cols = 0

        # Recognized block headings
        HEADING_KEYS = [
            "lesson information", "learning objectives", "target language",
            "lesson stages", "supporting details", "differentiation",
            "assessment and feedback", "assessment & feedback",
            "reflection and notes", "reflection & notes"
        ]

        # Label pattern for bulletizing supporting details
        LABEL_RE = re.compile(r"^([A-Z][A-Za-z &]+):\s*(.*)$")

        for raw in lesson_text.split("\n"):
            line = raw.strip()
            if not line:
                # Blank line ends supporting-details bullet mode
                if in_supporting_details:
                    in_supporting_details = False
                continue

            # SECTION 2 page break
            if "SECTION 2" in line.upper() and not inside_section2:
                close_table()
                doc.add_page_break()
                inside_section2 = True
                in_supporting_details = False
                continue

            # SECTION headers like "Section 1 — ..." or "SECTION 1 — ..."
            if re.match(r"^section\s+\d+", line, re.I):
                close_table()
                in_supporting_details = False
                p = doc.add_paragraph(line.upper())
                run = p.runs[0]
                run.font.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(255, 255, 255)
                shading = parse_xml(r'<w:shd {} w:fill="003366"/>'.format(nsdecls("w")))
                p._p.get_or_add_pPr().append(shading)
                p.alignment = 1
                doc.add_paragraph()
                continue

            # Pipe tables
            if "|" in line:
                cols = [c.strip() for c in line.split("|")]
                if current_table is None:
                    current_table = doc.add_table(rows=1, cols=len(cols))
                    current_table_cols = len(cols)
                    current_table.style = "Table Grid"
                    hdr_cells = current_table.rows[0].cells
                    for i, text in enumerate(cols):
                        hdr_cells[i].text = text
                        for pp in hdr_cells[i].paragraphs:
                            rr = pp.runs[0] if pp.runs else pp.add_run()
                            rr.font.bold = True
                            rr.font.size = Pt(10)
                    for cell in hdr_cells:
                        shading = parse_xml(r'<w:shd {} w:fill="E6E6FA"/>'.format(nsdecls("w")))
                        cell._tc.get_or_add_tcPr().append(shading)
                else:
                    if len(cols) < current_table_cols:
                        cols += [""] * (current_table_cols - len(cols))
                    elif len(cols) > current_table_cols:
                        cols = cols[:current_table_cols]
                    row = current_table.add_row()
                    for i, text in enumerate(cols):
                        row.cells[i].text = text
                continue

            # Domain table blocks in Section 2
            low = line.lower()
            if low.startswith("domain name"):
                close_table()
                in_supporting_details = False
                current_table = doc.add_table(rows=3, cols=2)
                current_table_cols = 2
                current_table.style = "Table Grid"
                for column in current_table.columns:
                    for cell in column.cells:
                        cell.width = Inches(3.5)
                hdr = current_table.rows[0].cells
                hdr[0].text = "Domain Name"
                hdr[1].text = re.sub(r"^domain name[:]*", "", line, flags=re.I).strip()
                hdr[0].paragraphs[0].runs[0].font.bold = True
                hdr[0]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls("w"))))
                continue

            if low.startswith("rubric check"):
                if current_table is not None and len(current_table.rows) >= 2:
                    row = current_table.rows[1]
                    row.cells[0].text = "Rubric Check"
                    row.cells[1].text = re.sub(r"^rubric check[:]*", "", line, flags=re.I).strip()
                    row.cells[0].paragraphs[0].runs[0].font.bold = True
                continue

            if low.startswith("ai mentor comment"):
                if current_table is not None and len(current_table.rows) >= 3:
                    row = current_table.rows[2]
                    row.cells[0].text = "AI Mentor Comment"
                    row.cells[1].text = re.sub(r"^ai mentor comment[:]*", "", line, flags=re.I).strip()
                    row.cells[0].paragraphs[0].runs[0].font.bold = True
                close_table()
                continue

            # Headings and heading+content on the same line
            # Detect any known heading at start, possibly followed by text.
            matched_heading = None
            for hk in HEADING_KEYS:
                if line.lower().startswith(hk):
                    matched_heading = hk
                    break

            if matched_heading:
                close_table()
                in_supporting_details = ("supporting details" in matched_heading)

                # Split heading from any trailing content on the same line
                trailing = line[len(matched_heading):].strip(" :—-")
                # Add heading as bold
                p = doc.add_paragraph(line[:len(matched_heading)])
                if p.runs:
                    p.runs[0].font.bold = True
                    p.runs[0].font.size = Pt(12)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(6)

                # If there is trailing content, add it as a normal paragraph (not bold)
                if trailing:
                    q = doc.add_paragraph(trailing)
                    q.paragraph_format.line_spacing = 1.15
                    q.paragraph_format.space_after = Pt(4)

                continue

            # Bullet list for Supporting Details
            if in_supporting_details:
                m = LABEL_RE.match(line)
                if m:
                    label, rest = m.group(1), m.group(2)
                    item = doc.add_paragraph(style=None)
                    item.style = doc.styles['List Bullet']
                    run_label = item.add_run(label + ": ")
                    run_label.font.bold = True
                    item.add_run(rest)
                else:
                    item = doc.add_paragraph(style=None)
                    item.style = doc.styles['List Bullet']
                    item.add_run(line)
                continue

            # Default paragraph
            close_table()
            in_supporting_details = False
            p = doc.add_paragraph(line)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(4)

        # Footer
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "AI Lesson Planner — BAE StanEval Hybrid | © 2025 Kaled Alenezi"
        footer_para.alignment = 1
        footer_para.runs[0].font.size = Pt(8)

        # Save and return
        output = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(output.name)
        output.seek(0)
        return send_file(output.name, as_attachment=True, download_name="BAE_Lesson_Plan.docx")

    except Exception as e:
        print("❌ ERROR in /generate:", e)
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    # If deploying on Railway/Gunicorn, this block is ignored (Gunicorn imports the app)
    app.run(host="0.0.0.0", port=5000)
